from __future__ import annotations

import json
import os
import re
import sys
import tempfile
import time
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any

from flask import Flask, after_this_request, jsonify, request, send_file

try:
    from openai import APIConnectionError, APIError, APITimeoutError, AuthenticationError, OpenAI, RateLimitError

    OPENAI_AVAILABLE = True
except ImportError:  # pragma: no cover - handled at runtime
    OPENAI_AVAILABLE = False
    OpenAI = None  # type: ignore[assignment]
    AuthenticationError = None  # type: ignore[assignment]
    RateLimitError = None  # type: ignore[assignment]
    APIConnectionError = None  # type: ignore[assignment]
    APITimeoutError = None  # type: ignore[assignment]
    APIError = None  # type: ignore[assignment]

try:
    from docxtpl import DocxTemplate, RichText
    from docx import Document
    from docx.enum.text import WD_COLOR_INDEX

    DOCX_AVAILABLE = True
except ImportError:  # pragma: no cover - handled at runtime
    DOCX_AVAILABLE = False
    DocxTemplate = None  # type: ignore[assignment]
    RichText = None  # type: ignore[assignment]
    Document = None  # type: ignore[assignment]
    WD_COLOR_INDEX = None  # type: ignore[assignment]


BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = BASE_DIR / "data"
DESCRIPTION_PATH = DATA_DIR / "description_merged.json"
SCHEMA_PATH = DATA_DIR / "output_schema.json"
TEMPLATE_PATH = DATA_DIR / "template.docx"


def load_env_file(path: Path, override: bool = True) -> None:
    if not path.exists():
        return

    for raw_line in path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#"):
            continue

        if "=" not in line:
            continue

        key, value = line.split("=", 1)
        key = key.strip()
        value = value.strip().strip('"').strip("'")
        if key and (override or key not in os.environ):
            os.environ[key] = value


load_env_file(BASE_DIR / ".env", override=True)
MODEL_NAME = os.getenv("OPENAI_MODEL", "gpt-4.1")
OPENAI_TIMEOUT_SECONDS = float(os.getenv("OPENAI_TIMEOUT_SECONDS", "90"))
OPENAI_MAX_RETRIES = int(os.getenv("OPENAI_MAX_RETRIES", "0"))

SYSTEM_PROMPT = """Ти досвідчений проджект-менеджер, який готує технічні завдання для сайтів.
Працюй за правилами:
1. Кожен новий запит повністю незалежний. Ігноруй будь-які попередні діалоги.
2. Аналізуй тільки поточний вхідний матеріал: бриф, опис, таблиці, текст, файли.
3. Не вигадуй нові дані. Використовуй лише факти з наданого матеріалу.
4. Жодна деталь не повинна бути пропущена: контакти, месенджери, адреси, функції, ролі, блоки.
5. Формулювання професійні, лаконічні, ділові, без емоцій і суб'єктивних оцінок.
5.1. Усі текстові описи формуй як готові формулювання для офіційного документа ТЗ.
5.2. Не скорочуй зміст у вигляді тез, якщо у вхідному матеріалі є детальний опис.
5.3. Не генеруй зайвий функціонал, який прямо не випливає з вхідних даних.
6. Якщо поле очікує масив і дані відсутні, поверни порожній масив [].
7. Якщо інформації для поля немає, заповни "уточнити".
8. Заповнюй JSON строго за заданою структурою.
9. Поверни тільки валідний JSON-об'єкт без пояснень і без markdown.
10. Для функціональних блоків (особливо admin_functions, user_functions, dynamic_sections):
10.1. Формуй опис у професійному стилі ТЗ: 1-2 повні речення, без телеграфних тез.
10.2. Обов'язково вкажи суть функції, ключові дії користувача/адміністратора та очікуваний результат.
10.3. Не спрощуй до коротких формулювань на кшталт "Імпорт товарів" або "Пошук".
10.4. Зберігай нейтральну ділову лексику та однаковий стиль у всіх пунктах.
"""


def load_json_file(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {}

    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def build_default_from_schema(schema: Any) -> Any:
    if isinstance(schema, dict):
        return {key: build_default_from_schema(value) for key, value in schema.items()}

    if isinstance(schema, list):
        return []

    return "уточнити"


def normalize_by_schema(value: Any, schema: Any) -> Any:
    if isinstance(schema, dict):
        source = value if isinstance(value, dict) else {}
        result: dict[str, Any] = {}
        for key, sub_schema in schema.items():
            result[key] = normalize_by_schema(source.get(key), sub_schema)
        return result

    if isinstance(schema, list):
        if not isinstance(value, list):
            return []

        if not schema:
            return value

        item_schema = schema[0]
        return [normalize_by_schema(item, item_schema) for item in value]

    if value is None:
        return "уточнити"

    if isinstance(value, str):
        stripped = value.strip()
        return stripped if stripped else "уточнити"

    return value


def is_placeholder_text(value: Any) -> bool:
    if not isinstance(value, str):
        return False
    normalized = value.strip().lower()
    return normalized in {"", "уточнити", "n/a", "none", "-"}


def has_meaningful_value(value: Any) -> bool:
    if value is None:
        return False
    if isinstance(value, str):
        return bool(value.strip())
    if isinstance(value, list):
        return any(has_meaningful_value(item) for item in value)
    if isinstance(value, dict):
        return any(has_meaningful_value(item) for item in value.values())
    return True


def prune_payload_lists(value: Any) -> Any:
    if isinstance(value, dict):
        return {key: prune_payload_lists(item) for key, item in value.items()}

    if isinstance(value, list):
        cleaned: list[Any] = []
        for item in value:
            cleaned_item = prune_payload_lists(item)
            if has_meaningful_value(cleaned_item):
                cleaned.append(cleaned_item)
        deduplicated: list[Any] = []
        seen: set[str] = set()
        for item in cleaned:
            try:
                key = json.dumps(item, ensure_ascii=False, sort_keys=True)
            except TypeError:
                key = str(item)
            if key in seen:
                continue
            seen.add(key)
            deduplicated.append(item)
        return deduplicated

    if isinstance(value, str):
        return value.strip()

    return value


def auto_fill_pages_count(payload: dict[str, Any]) -> dict[str, Any]:
    current = payload.get("pages_count")
    if isinstance(current, str) and current.strip() and current.strip().lower() != "уточнити":
        return payload

    sections = payload.get("site_sections")
    if not isinstance(sections, list):
        payload["pages_count"] = "уточнити"
        return payload

    normalized_sections = []
    for section in sections:
        if isinstance(section, str):
            name = section.strip()
            if name:
                normalized_sections.append(name)

    unique_sections: list[str] = []
    seen: set[str] = set()
    for section in normalized_sections:
        key = section.lower()
        if key not in seen:
            seen.add(key)
            unique_sections.append(section)

    count = len(unique_sections)
    if normalize_ecommerce_flag(payload.get("is_ecommerce")) and not any("каталог" in s.lower() for s in unique_sections):
        count += 1

    payload["pages_count"] = str(count) if count > 0 else "уточнити"
    payload["site_sections"] = unique_sections
    return payload


def sanitize_template_fields(payload: dict[str, Any]) -> dict[str, Any]:
    customer_name = payload.get("customer_name")
    if isinstance(customer_name, str) and customer_name.strip() and not is_placeholder_text(customer_name):
        payload["main_products"] = customer_name.strip()

    if is_placeholder_text(payload.get("domain")):
        payload["domain"] = ""

    if isinstance(payload.get("competitors_list"), list) and not payload["competitors_list"]:
        payload["competitors_list"] = ["уточнити (надати посилання на сайти конкурентів)"]

    for key in ("footer_address_note", "footer_payment_note", "footer_notes", "product_card_note", "checkout_tech_note"):
        if is_placeholder_text(payload.get(key)):
            payload[key] = ""

    header_blocks = payload.get("header_blocks")
    if isinstance(header_blocks, list):
        for block in header_blocks:
            if not isinstance(block, dict):
                continue
            if is_placeholder_text(block.get("note")):
                block["note"] = ""
            subitems = block.get("subitems")
            if isinstance(subitems, list):
                block["subitems"] = [s for s in subitems if isinstance(s, str) and not is_placeholder_text(s)]

    main_blocks = payload.get("main_page_blocks")
    if isinstance(main_blocks, list):
        for block in main_blocks:
            if not isinstance(block, dict):
                continue
            for key in ("client_note", "visual_note", "note_for_comp"):
                if is_placeholder_text(block.get(key)):
                    block[key] = ""

    dynamic_sections = payload.get("dynamic_sections")
    if isinstance(dynamic_sections, list):
        for section in dynamic_sections:
            if not isinstance(section, dict):
                continue
            if is_placeholder_text(section.get("note_for_comp")):
                section["note_for_comp"] = ""

    return payload


ADMIN_TOPIC_KEYWORDS = (
    "адмін",
    "адміністра",
    "управл",
    "керув",
    "налаштув",
    "доступ",
    "рол",
    "права",
    "api",
    "інтеграц",
    "синхрон",
    "імпорт",
    "експорт",
    "модерац",
    "журнал",
    "лог",
    "звіт",
    "аналіт",
    "бекап",
    "резерв",
    "бекофіс",
    "backoffice",
    "cms",
    "облік",
    "склад",
    "логіст",
    "workflow",
    "штрихкод",
    "сканув",
    "erp",
    "crm",
    "acl",
    "permission",
)

DEFAULT_USER_SECTIONS = [
    "Головна",
    "Про компанію",
    "Доставка та оплата",
    "Контакти",
    "Реєстрація/Вхід",
    "Особистий кабінет",
]


def is_admin_topic(text: str) -> bool:
    lower = text.lower()
    return any(keyword in lower for keyword in ADMIN_TOPIC_KEYWORDS)


def split_sentences(text: str) -> list[str]:
    raw_parts = re.split(r"[.\n]+", text)
    return [part.strip(" -\t\r") for part in raw_parts if part.strip(" -\t\r")]


def detect_admin_category(source_text: str) -> str:
    text = source_text.lower()
    has_role_keyword = re.search(r"\bрол(ь|і|ей|ями|ях)?\b", text) is not None
    if has_role_keyword or any(token in text for token in ("права доступу", "доступ", "permission", "acl", "auth")):
        return "access"
    if "імпорт" in text:
        return "import"
    if "експорт" in text:
        return "export"
    if any(token in text for token in ("скан", "qr", "barcode", "штрихкод", "ідентифікатор")):
        return "scan"
    if any(token in text for token in ("пошук", "фільтр", "filter", "search")):
        return "search"
    if any(token in text for token in ("звіт", "аналіт", "dashboard", "report")):
        return "analytics"
    if any(token in text for token in ("замовлен", "order")):
        return "orders"
    if any(token in text for token in ("користувач", "user", "акаунт", "обліков")):
        return "users"
    if any(token in text for token in ("api", "інтеграц", "синхрон", "erp", "crm", "webhook")):
        return "integration"
    if any(token in text for token in ("облік", "склад", "логіст", "залишк")):
        return "inventory"
    if any(token in text for token in ("управ", "керув", "редаг", "додав", "видал", "налаштув", "crud")):
        return "management"
    return "generic"


def build_admin_function_name(sentence: str) -> str:
    category = detect_admin_category(sentence)
    category_names = {
        "access": "Ролі та права доступу",
        "import": "Імпорт даних",
        "export": "Експорт даних",
        "scan": "Сканування ідентифікаторів",
        "search": "Розширений пошук",
        "analytics": "Звіти та аналітика",
        "orders": "Керування замовленнями",
        "users": "Керування користувачами",
        "integration": "Інтеграція із зовнішніми системами",
        "inventory": "Облік і контроль ресурсів",
        "management": "Керування даними",
    }
    if category in category_names:
        return category_names[category]
    words = sentence.split()
    return " ".join(words[:5]) if words else "уточнити"


USER_SECTION_TEMPLATES: list[tuple[list[str], dict[str, Any]]] = [
    (
        ["головна"],
        {
            "description": "Сторінка першого контакту з користувачем: презентація пропозиції, ключові переваги, переходи у каталог та основні розділи.",
            "fields": [
                {"name": "Перший екран", "description": "Заголовок, короткий опис, ключова дія."},
                {"name": "Переваги", "description": "Стислий перелік переваг компанії."},
                {"name": "Навігація в каталог", "description": "Швидкі переходи у популярні категорії."},
            ],
        },
    ),
    (
        ["про компан", "about"],
        {
            "description": "Інформаційна сторінка про компанію, її спеціалізацію, досвід та цінності бренду.",
            "fields": [
                {"name": "Опис компанії", "description": "Коротка та розширена презентація бренду."},
                {"name": "Переваги компанії", "description": "Аргументи довіри для клієнта."},
                {"name": "Контактні дані", "description": "Актуальні канали зв'язку."},
            ],
        },
    ),
    (
        ["доставка", "оплата", "shipping", "payment"],
        {
            "description": "Сторінка з умовами оплати, доставки, термінами обробки та вартістю сервісів.",
            "fields": [
                {"name": "Способи доставки", "description": "Доступні служби та терміни доставки."},
                {"name": "Способи оплати", "description": "Умови оплати та підтвердження платежу."},
                {"name": "Обмеження та винятки", "description": "Критичні умови обслуговування."},
            ],
        },
    ),
    (
        ["повернен", "refund", "return"],
        {
            "description": "Сторінка з політикою повернення, обміну товару та порядком обробки звернень.",
            "fields": [
                {"name": "Умови повернення", "description": "Підстави, терміни та обмеження."},
                {"name": "Процедура повернення", "description": "Покроковий порядок дій для клієнта."},
                {"name": "Компенсація", "description": "Правила відшкодування коштів."},
            ],
        },
    ),
    (
        ["контакт"],
        {
            "description": "Сторінка контактної інформації з основними каналами зв'язку та формою звернення.",
            "fields": [
                {"name": "Контакти", "description": "Телефон, email, месенджери."},
                {"name": "Адреса", "description": "Локація або реквізити для зв'язку."},
                {"name": "Форма звернення", "description": "Поля для надсилання повідомлення."},
            ],
        },
    ),
    (
        ["реєстрац", "вхід", "login", "sign"],
        {
            "description": "Сторінка авторизації та реєстрації для доступу до персональних функцій користувача.",
            "fields": [
                {"name": "Реєстрація", "description": "Створення нового облікового запису."},
                {"name": "Вхід", "description": "Авторизація існуючого користувача."},
                {"name": "Відновлення доступу", "description": "Процедура відновлення пароля."},
            ],
        },
    ),
    (
        ["кабінет", "profile", "account"],
        {
            "description": "Персональний кабінет користувача для керування профілем, замовленнями та збереженими даними.",
            "fields": [
                {"name": "Профіль", "description": "Редагування персональних даних."},
                {"name": "Історія замовлень", "description": "Перегляд статусів і деталей покупок."},
                {"name": "Збережені дані", "description": "Адреси доставки, обрані товари та налаштування."},
            ],
        },
    ),
]


def build_user_section_template(title: str) -> dict[str, Any]:
    lower = title.lower()
    for keywords, template in USER_SECTION_TEMPLATES:
        if any(keyword in lower for keyword in keywords):
            return {
                "title": title,
                "description": template["description"],
                "note_for_comp": "",
                "fields": template["fields"],
            }
    return {
        "title": title,
        "description": "Опис розділу необхідно деталізувати відповідно до погодженого функціоналу та контент-плану.",
        "note_for_comp": "",
        "fields": [
            {"name": "Контент", "description": "Перелік інформації, яка має бути опублікована у розділі."},
            {"name": "Функціонал", "description": "Опис дій користувача та реакції системи у межах розділу."},
        ],
    }


def repartition_admin_and_user_blocks(payload: dict[str, Any], material_text: str) -> dict[str, Any]:
    admin_functions = payload.get("admin_functions") if isinstance(payload.get("admin_functions"), list) else []
    dynamic_sections = payload.get("dynamic_sections") if isinstance(payload.get("dynamic_sections"), list) else []
    site_sections = payload.get("site_sections") if isinstance(payload.get("site_sections"), list) else []

    filtered_site_sections: list[str] = []
    for section in site_sections:
        if not isinstance(section, str):
            continue
        name = section.strip()
        if not name:
            continue
        if is_admin_topic(name):
            admin_functions.append({"name": build_admin_function_name(name), "description": name})
            continue
        filtered_site_sections.append(name)
    payload["site_sections"] = filtered_site_sections

    kept_sections: list[dict[str, Any]] = []
    for section in dynamic_sections:
        if not isinstance(section, dict):
            continue
        title = str(section.get("title", "")).strip()
        description = str(section.get("description", "")).strip()
        joined = f"{title} {description}".strip()
        if joined and is_admin_topic(joined):
            admin_functions.append(
                {
                    "name": build_admin_function_name(title or description),
                    "description": description or "уточнити",
                }
            )
            continue
        kept_sections.append(section)

    for sentence in split_sentences(material_text):
        lower_sentence = sentence.lower()
        if lower_sentence.startswith(("сторінки", "site sections", "sections", "pages")):
            continue
        if is_admin_topic(sentence):
            admin_functions.append(
                {
                    "name": build_admin_function_name(sentence),
                    "description": sentence,
                }
            )

    unique_admin_by_name: dict[str, dict[str, str]] = {}
    for item in admin_functions:
        if not isinstance(item, dict):
            continue
        name = str(item.get("name", "")).strip() or "уточнити"
        description = str(item.get("description", "")).strip() or "уточнити"
        name_key = name.lower()
        existing = unique_admin_by_name.get(name_key)
        if not existing:
            unique_admin_by_name[name_key] = {"name": name, "description": description}
            continue
        if len(description) > len(existing["description"]):
            unique_admin_by_name[name_key] = {"name": name, "description": description}

    unique_admin = list(unique_admin_by_name.values())
    payload["admin_functions"] = unique_admin or [{"name": "уточнити", "description": "уточнити"}]
    payload["dynamic_sections"] = kept_sections
    return payload


def enrich_dynamic_sections(payload: dict[str, Any]) -> dict[str, Any]:
    source_sections = payload.get("site_sections") if isinstance(payload.get("site_sections"), list) else []
    existing_sections = payload.get("dynamic_sections") if isinstance(payload.get("dynamic_sections"), list) else []

    existing_map: dict[str, dict[str, Any]] = {}
    for section in existing_sections:
        if not isinstance(section, dict):
            continue
        title = str(section.get("title", "")).strip()
        if not title:
            continue
        existing_map[title.lower()] = section

    candidate_sections: list[str] = []
    for section in source_sections:
        if not isinstance(section, str):
            continue
        title = section.strip()
        if not title:
            continue
        lower_title = title.lower()
        if "каталог" in lower_title:
            continue
        if is_admin_topic(title):
            continue
        candidate_sections.append(title)

    if not candidate_sections:
        for section in existing_sections:
            if not isinstance(section, dict):
                continue
            title = str(section.get("title", "")).strip()
            if not title:
                continue
            lower_title = title.lower()
            if "каталог" in lower_title:
                continue
            if is_admin_topic(title):
                continue
            candidate_sections.append(title)

    if not candidate_sections:
        candidate_sections = DEFAULT_USER_SECTIONS.copy()

    final_sections: list[dict[str, Any]] = []
    for title in candidate_sections:
        existing = existing_map.get(title.lower())
        if existing:
            description = str(existing.get("description", "")).strip()
            fields = existing.get("fields") if isinstance(existing.get("fields"), list) else []
            meaningful_fields = [f for f in fields if isinstance(f, dict) and has_meaningful_value(f)]
            if description and not is_placeholder_text(description) and meaningful_fields:
                final_sections.append(existing)
                continue

        final_sections.append(build_user_section_template(title))

    payload["dynamic_sections"] = final_sections or [build_user_section_template("уточнити")]
    return payload


def ensure_sentence(text: str) -> str:
    cleaned = text.strip()
    if not cleaned:
        return cleaned
    if cleaned[-1] in ".!?":
        return cleaned
    return f"{cleaned}."


def has_detailed_tz_description(value: str) -> bool:
    cleaned = value.strip()
    if not cleaned or is_placeholder_text(cleaned):
        return False
    return len(cleaned) >= 90 and len(cleaned.split()) >= 12


def infer_entity_key(source_text: str) -> str:
    text = source_text.lower()
    mapping: list[tuple[tuple[str, ...], str]] = [
        (("товар", "продукт", "номенклатур"), "product"),
        (("замовлен", "order"), "order"),
        (("курс", "урок", "модул", "навчан"), "course"),
        (("студент", "учень", "слухач"), "student"),
        (("користувач", "user", "акаунт", "обліков"), "user"),
        (("контент", "сторінк", "публікац", "матеріал"), "content"),
        (("файл", "документ"), "file"),
        (("категор", "довідник", "класифікатор"), "directory"),
        (("плат", "оплат"), "payment"),
        (("доставк", "відправ"), "delivery"),
        (("залишк", "склад", "облік", "ресурс"), "inventory"),
    ]
    for keywords, key in mapping:
        if any(keyword in text for keyword in keywords):
            return key
    return "generic"


def get_entity_forms(entity_key: str) -> dict[str, str]:
    forms: dict[str, dict[str, str]] = {
        "product": {
            "gen": "товарних позицій",
            "instr": "товарними позиціями",
        },
        "order": {
            "gen": "замовлень",
            "instr": "замовленнями",
        },
        "course": {
            "gen": "навчальних курсів",
            "instr": "навчальними курсами",
        },
        "student": {
            "gen": "даних студентів",
            "instr": "даними студентів",
        },
        "user": {
            "gen": "облікових записів користувачів",
            "instr": "обліковими записами користувачів",
        },
        "content": {
            "gen": "контентних матеріалів",
            "instr": "контентними матеріалами",
        },
        "file": {
            "gen": "файлів та документів",
            "instr": "файлами та документами",
        },
        "directory": {
            "gen": "довідників та класифікаторів",
            "instr": "довідниками та класифікаторами",
        },
        "payment": {
            "gen": "платіжних даних",
            "instr": "платіжними даними",
        },
        "delivery": {
            "gen": "даних доставки",
            "instr": "даними доставки",
        },
        "inventory": {
            "gen": "облікових даних і ресурсів",
            "instr": "обліковими даними та ресурсами",
        },
        "generic": {
            "gen": "даних системи",
            "instr": "даними системи",
        },
    }
    return forms.get(entity_key, forms["generic"])


def infer_external_system_name(source_text: str, material_text: str = "") -> str | None:
    combined = f"{source_text}\n{material_text}".strip()
    if not combined:
        return None

    domain_match = re.search(r"https?://([A-Za-z0-9.-]+\.[A-Za-z]{2,})", combined)
    if domain_match:
        return domain_match.group(1)

    skip = {"API", "CRM", "ERP", "CMS", "SQL", "JSON", "REST", "SOAP", "HTTP", "HTTPS"}
    for match in re.finditer(r"(?:з|до|через|via)\s+([A-Za-z][A-Za-z0-9_.-]{2,})", combined, flags=re.IGNORECASE):
        candidate = match.group(1)
        if candidate.upper() not in skip:
            return candidate

    for match in re.finditer(r"\b([A-Z][A-Za-z0-9_.-]{2,})\b", combined):
        candidate = match.group(1)
        if candidate.upper() not in skip:
            return candidate
    return None


def build_admin_description(name: str, description: str, material_text: str) -> str:
    raw_source = f"{name} {description}".strip()
    source = raw_source.lower()
    material_lower = material_text.lower()
    entity_key = infer_entity_key(raw_source)
    entity = get_entity_forms(entity_key)
    category = detect_admin_category(raw_source)

    if category == "management":
        return (
            f"Функціонал повинен забезпечувати централізоване керування {entity['instr']} в адміністративній панелі: "
            "створення, редагування, архівацію або видалення записів, а також контроль актуальності пов'язаних параметрів "
            "відповідно до затверджених бізнес-правил проєкту."
        )
    if category == "import":
        base = (
            f"Система повинна підтримувати масове завантаження {entity['gen']} через файл або зовнішнє джерело з автоматичною "
            "перевіркою структури даних, обробкою помилок та коректним оновленням існуючих записів."
        )
        formats: list[str] = []
        if "excel" in source or "excel" in material_lower or "xlsx" in source or "xls" in source:
            formats.append("Excel")
        if "csv" in source or "csv" in material_lower:
            formats.append("CSV")
        if "xml" in source or "xml" in material_lower:
            formats.append("XML")
        if "json" in source or "json" in material_lower:
            formats.append("JSON")
        if formats:
            return (
                base
                + f" Підтримувані формати: {', '.join(formats)}; склад полів та правила відповідності мають відповідати погодженому шаблону."
            )
        return base
    if category == "export":
        return (
            f"Необхідно реалізувати вивантаження {entity['gen']} у погоджені формати з налаштуванням складу полів, "
            "контролем цілісності даних та можливістю подальшого використання у зовнішніх процесах."
        )
    if category == "search":
        return (
            f"Передбачити розширений пошук і фільтрацію {entity['gen']} за ключовими атрибутами та службовими параметрами, "
            "щоб забезпечити швидкий доступ до потрібних записів і скоротити час операційної обробки."
        )
    if category == "scan":
        return (
            "Після сканування ідентифікатора система повинна виконувати автоматичну перевірку наявності запису в базі: "
            "відкривати знайдений елемент для подальших дій або запускати сценарій створення нового запису з попередньо заповненими полями."
        )
    if "штрихкод" in source:
        return (
            "Функціонал має забезпечувати роботу зі штрихкодами для швидкої ідентифікації об'єктів, "
            "пошуку записів у системі та зниження кількості помилок при операційній обробці."
        )
    if category == "orders":
        return (
            "Адміністративна частина повинна забезпечувати повний цикл керування замовленнями: перегляд складу, "
            "зміну статусів обробки, фіксацію ключових етапів виконання та контроль актуального стану по кожній заявці."
        )
    if category == "users":
        return (
            "Необхідно реалізувати керування обліковими записами користувачів: створення, редагування, блокування "
            "та контроль актуальності даних у межах затвердженої моделі доступу."
        )
    if category == "access":
        return (
            "Система повинна підтримувати рольову модель доступу: призначення ролей, налаштування прав та контроль "
            "доступу до адміністративних функцій відповідно до політики безпеки проєкту."
        )
    if category == "integration":
        system_name = infer_external_system_name(raw_source, material_text)
        system_part = f" із зовнішньою системою {system_name}" if system_name else " із зовнішніми системами"
        return (
            f"Реалізувати інтеграцію{system_part} через API з визначенням напрямків обміну, "
            "переліку сутностей, періодичності синхронізації та правил обробки помилок для стабільної роботи процесів."
        )
    if category == "inventory":
        return (
            "Функціонал повинен підтримувати облік ресурсів і операцій в межах адміністративної панелі, "
            "щоб забезпечити актуальність даних та керованість внутрішніх процесів."
        )
    if category == "analytics":
        return (
            "Необхідно реалізувати формування звітів та аналітики за ключовими показниками процесів, "
            "щоб забезпечити контроль результатів, прозорість даних і підтримку управлінських рішень."
        )

    if description and not is_placeholder_text(description):
        return (
            f"{ensure_sentence(description)} "
            "Функціонал реалізується в адміністративній панелі та деталізується в межах погоджених бізнес-правил проєкту."
        )
    return "уточнити"


def admin_semantic_key(name: str, description: str) -> str:
    text = f"{name} {description}".strip()
    category = detect_admin_category(text)
    if category == "integration":
        system_name = infer_external_system_name(text) or "generic"
        return f"integration:{system_name.lower()}"
    if category in {"management", "import", "export", "search", "inventory"}:
        name_entity = infer_entity_key(name)
        return f"{category}:{name_entity}"
    if category != "generic":
        return category
    return name.strip().lower() or "уточнити"


def enrich_admin_functions(payload: dict[str, Any], material_text: str) -> dict[str, Any]:
    raw = payload.get("admin_functions")
    if not isinstance(raw, list):
        payload["admin_functions"] = [{"name": "уточнити", "description": "уточнити"}]
        return payload

    enriched: list[dict[str, str]] = []
    for item in raw:
        if not isinstance(item, dict):
            continue
        name = str(item.get("name", "")).strip()
        description = str(item.get("description", "")).strip()

        if not name and description and not is_placeholder_text(description):
            name = build_admin_function_name(description)
        if not name:
            name = "уточнити"

        if has_detailed_tz_description(description):
            final_description = ensure_sentence(description)
        else:
            final_description = build_admin_description(name, description, material_text)

        enriched.append({"name": name, "description": final_description})

    deduped: dict[str, dict[str, str]] = {}
    for item in enriched:
        key = admin_semantic_key(item["name"], item["description"])
        existing = deduped.get(key)
        if not existing:
            deduped[key] = item
            continue
        prefer_current = False
        if len(item["description"]) > len(existing["description"]):
            prefer_current = True
        if key.startswith("integration:") and key != "integration:generic":
            prefer_current = True
        if prefer_current:
            deduped[key] = item

    cleaned_keys = set(deduped.keys())
    for key in list(cleaned_keys):
        if not key.endswith(":generic"):
            continue
        prefix = key.rsplit(":", 1)[0]
        has_specific = any(other.startswith(f"{prefix}:") and other != key for other in cleaned_keys)
        if has_specific:
            deduped.pop(key, None)

    payload["admin_functions"] = list(deduped.values()) or [{"name": "уточнити", "description": "уточнити"}]
    return payload


def enrich_main_page_blocks(payload: dict[str, Any]) -> dict[str, Any]:
    blocks = payload.get("main_page_blocks")
    if not isinstance(blocks, list):
        blocks = []

    normalized: list[dict[str, Any]] = [item for item in blocks if isinstance(item, dict)]
    names = {str(item.get("name", "")).strip().lower() for item in normalized}

    defaults = [
        {
            "name": "Перший екран",
            "description": "Ключова пропозиція компанії, короткий опис спеціалізації та основна цільова дія для користувача.",
            "client_note": "",
            "visual_note": "Заголовок, підзаголовок, кнопка дії, фонове зображення або графічний акцент.",
            "note_for_comp": "Уточнити фінальний текст УТП та CTA.",
        },
        {
            "name": "Категорії товарів",
            "description": "Огляд основних товарних категорій з переходом у відповідні розділи каталогу.",
            "client_note": "",
            "visual_note": "Сітка карток категорій з зображенням, назвою та посиланням.",
            "note_for_comp": "Уточнити перелік категорій та порядок пріоритету.",
        },
        {
            "name": "Хіти та новинки",
            "description": "Блок промо-товарів для швидкого ознайомлення з актуальним асортиментом.",
            "client_note": "",
            "visual_note": "Карусель або сітка товарів з ціною, кнопкою переходу та короткими характеристиками.",
            "note_for_comp": "Уточнити джерело відбору товарів (хіти/новинки/акції).",
        },
        {
            "name": "Переваги компанії",
            "description": "Аргументація цінності бренду для формування довіри та підвищення конверсії.",
            "client_note": "",
            "visual_note": "Іконки або інфоблоки з короткими тезами та поясненням.",
            "note_for_comp": "Уточнити затверджений список переваг.",
        },
        {
            "name": "Оплата, доставка, повернення",
            "description": "Короткі правила оформлення, оплати, доставки та повернення для покупця.",
            "client_note": "",
            "visual_note": "Структуровані картки або акордеон з ключовими умовами.",
            "note_for_comp": "Уточнити фінальні бізнес-правила для публікації.",
        },
        {
            "name": "Контактний блок",
            "description": "Контакти компанії та форма звернення для оперативного зв'язку з клієнтом.",
            "client_note": "",
            "visual_note": "Телефон, email, месенджери, адреса, кнопки швидкого зв'язку.",
            "note_for_comp": "Уточнити актуальні контактні дані та відповідального менеджера.",
        },
    ]

    if normalized:
        normalized = [
            item
            for item in normalized
            if not (is_placeholder_text(item.get("name")) and is_placeholder_text(item.get("description")))
        ]
        names = {str(item.get("name", "")).strip().lower() for item in normalized}

    for item in defaults:
        key = item["name"].strip().lower()
        if key in names:
            continue
        normalized.append(item)
        names.add(key)
        if len(normalized) >= 6:
            break

    payload["main_page_blocks"] = normalized if normalized else defaults
    return payload


def merge_missing_from_enriched(payload: dict[str, Any], enriched: dict[str, Any]) -> dict[str, Any]:
    for key, value in enriched.items():
        current = payload.get(key)

        if isinstance(current, str):
            if is_placeholder_text(current) and isinstance(value, str) and not is_placeholder_text(value):
                payload[key] = value
            continue

        if isinstance(current, list):
            if not current and isinstance(value, list) and value:
                payload[key] = value
            continue

        if current is None and value is not None:
            payload[key] = value

    return payload


def ensure_required_template_blocks(payload: dict[str, Any]) -> dict[str, Any]:
    def ensure_list_of_dicts(key: str, template: dict[str, Any]) -> None:
        arr = payload.get(key)
        if isinstance(arr, list) and arr:
            return
        payload[key] = [template]

    def ensure_list_of_strings(key: str, fallback: str = "уточнити") -> None:
        arr = payload.get(key)
        if isinstance(arr, list) and arr:
            return
        payload[key] = [fallback]

    ensure_list_of_strings("site_sections")
    ensure_list_of_dicts("user_roles", {"name": "уточнити", "description": "уточнити"})
    ensure_list_of_dicts("admin_functions", {"name": "уточнити", "description": "уточнити"})
    ensure_list_of_dicts("user_functions", {"name": "уточнити", "description": "уточнити"})
    ensure_list_of_dicts("target_audience_groups", {"name": "уточнити", "description": "уточнити"})
    ensure_list_of_dicts(
        "header_blocks",
        {"title": "уточнити", "description": "уточнити", "subitems": [], "note": ""},
    )
    ensure_list_of_dicts(
        "main_page_blocks",
        {
            "name": "уточнити",
            "description": "уточнити",
            "client_note": "",
            "visual_note": "",
            "note_for_comp": "",
        },
    )

    ecommerce = normalize_ecommerce_flag(payload.get("is_ecommerce"))
    if ecommerce:
        ensure_list_of_dicts(
            "product_admin_fields",
            {
                "name": "уточнити",
                "description": "уточнити",
                "subfields": [{"name": "уточнити", "description": "уточнити"}],
            },
        )
        ensure_list_of_strings("cart_columns")
        ensure_list_of_strings("checkout_registered_fields")
        ensure_list_of_strings("checkout_guest_fields")

    dynamic_sections = payload.get("dynamic_sections")
    if not isinstance(dynamic_sections, list) or not dynamic_sections:
        source_sections = payload.get("site_sections") if isinstance(payload.get("site_sections"), list) else []
        generated_sections: list[dict[str, Any]] = []
        for section in source_sections:
            if not isinstance(section, str):
                continue
            title = section.strip()
            if not title:
                continue
            if "каталог" in title.lower():
                continue
            generated_sections.append(
                {
                    "title": title,
                    "description": "уточнити",
                    "note_for_comp": "",
                    "fields": [{"name": "уточнити", "description": "уточнити"}],
                }
            )
        payload["dynamic_sections"] = generated_sections or [
            {
                "title": "уточнити",
                "description": "уточнити",
                "note_for_comp": "",
                "fields": [{"name": "уточнити", "description": "уточнити"}],
            }
        ]

    return payload


def finalize_payload_for_template(payload: dict[str, Any], material_text: str, schema: dict[str, Any]) -> dict[str, Any]:
    cleaned = prune_payload_lists(payload)
    if isinstance(cleaned, dict):
        enriched = enrich_from_text(material_text, schema)
        cleaned = merge_missing_from_enriched(cleaned, enriched)
        cleaned = repartition_admin_and_user_blocks(cleaned, material_text)
        cleaned = enrich_admin_functions(cleaned, material_text)
        cleaned = auto_fill_pages_count(cleaned)
        cleaned = ensure_required_template_blocks(cleaned)
        cleaned = enrich_dynamic_sections(cleaned)
        cleaned = enrich_main_page_blocks(cleaned)
        cleaned = sanitize_template_fields(cleaned)
    return cleaned


def strip_markdown_fence(text: str) -> str:
    cleaned = text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\\s*", "", cleaned)
        cleaned = re.sub(r"\\s*```$", "", cleaned)
    return cleaned.strip()


def parse_json_from_response(text: str) -> dict[str, Any]:
    cleaned = strip_markdown_fence(text)
    return json.loads(cleaned)


def safe_error_message(exc: Exception) -> str:
    if OPENAI_AVAILABLE and AuthenticationError and isinstance(exc, AuthenticationError):
        return "invalid_api_key"
    if OPENAI_AVAILABLE and RateLimitError and isinstance(exc, RateLimitError):
        return "rate_limited"
    if OPENAI_AVAILABLE and APITimeoutError and isinstance(exc, APITimeoutError):
        return "openai_timeout"
    if OPENAI_AVAILABLE and APIConnectionError and isinstance(exc, APIConnectionError):
        return "connection_error"
    if OPENAI_AVAILABLE and APIError and isinstance(exc, APIError):
        return "openai_api_error"
    message = str(exc).strip()
    if not message:
        return exc.__class__.__name__
    return message[:180]


def read_uploaded_text(file_obj: Any) -> str:
    if not file_obj:
        return ""

    data = file_obj.read()
    if not data:
        return ""

    for encoding in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue

    return data.decode("utf-8", errors="ignore")


def get_generation_payload(material_text: str, placeholders: dict[str, Any], schema: dict[str, Any]) -> str:
    return (
        "Сформуй повний JSON для технічного завдання на основі вхідного матеріалу.\\n\\n"
        f"Опис плейсхолдерів (description_merged.json):\\n{json.dumps(placeholders, ensure_ascii=False, indent=2)}\\n\\n"
        f"Цільова структура JSON (дотримуйся ключів і вкладеності):\\n{json.dumps(schema, ensure_ascii=False, indent=2)}\\n\\n"
        f"Вхідний матеріал для аналізу:\\n{material_text}"
    )


def split_list_values(raw_value: str) -> list[str]:
    value = raw_value.strip()
    if not value:
        return []

    numbered_parts = re.split(r"(?=\d+(?:\.\d+)*\.\s+)", value)
    fragments: list[str] = []
    for part in numbered_parts:
        cleaned_part = re.sub(r"^\d+(?:\.\d+)*\.\s*", "", part).strip()
        if not cleaned_part:
            continue
        fragments.extend(re.split(r"[,;|\n]+", cleaned_part))

    normalized: list[str] = []
    seen: set[str] = set()
    for part in fragments:
        cleaned = part.strip(" -\t\r")
        if not cleaned:
            continue
        key = cleaned.lower()
        if key in seen:
            continue
        seen.add(key)
        normalized.append(cleaned)
    return normalized


def extract_by_patterns(material_text: str, patterns: list[str]) -> str | None:
    for pattern in patterns:
        match = re.search(pattern, material_text, flags=re.IGNORECASE)
        if match:
            value = match.group(1).strip()
            if value:
                return value
    return None


def enrich_from_text(material_text: str, schema: dict[str, Any]) -> dict[str, Any]:
    result = build_default_from_schema(schema)

    customer_name = extract_by_patterns(
        material_text,
        [
            r"(?:компанія|назва компанії|замовник|бренд)\s*[:\-]\s*(.+)",
            r"(?:company|client|customer)\s*[:\-]\s*(.+)",
        ],
    )
    if customer_name:
        result["customer_name"] = customer_name

    domain_match = re.search(r"\b(?:https?://)?([a-z0-9][a-z0-9.-]+\.[a-z]{2,})(?:/\S*)?\b", material_text, re.I)
    if domain_match:
        result["domain"] = domain_match.group(1).lower()

    email_match = re.search(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", material_text, re.I)
    if email_match:
        result["email"] = email_match.group(0)

    phones = re.findall(r"(?:\+?\d[\d()\-\s]{7,}\d)", material_text)
    clean_phones = [re.sub(r"\s+", " ", phone).strip() for phone in phones]
    if clean_phones:
        result["footer_phone_list"] = clean_phones

    languages_line = extract_by_patterns(
        material_text,
        [r"(?:мови|мова|languages?)\s*[:\-]\s*(.+)"],
    )
    if languages_line:
        result["languages"] = split_list_values(languages_line)

    sections_line = extract_by_patterns(
        material_text,
        [
            r"(?:розділи сайту|сторінки сайту|сторінки|sections?|pages?)\s*[:\-]\s*(.+)",
            r"(?:необхідні розділи|site sections)\s*[:\-]\s*(.+)",
        ],
    )
    if sections_line:
        result["site_sections"] = split_list_values(sections_line)

    schedule = extract_by_patterns(
        material_text,
        [r"(?:графік роботи|режим роботи|schedule|working hours)\s*[:\-]\s*(.+)"],
    )
    if schedule:
        result["schedule"] = schedule

    address = extract_by_patterns(
        material_text,
        [r"(?:адреса|address)\s*[:\-]\s*(.+)"],
    )
    if address:
        result["footer_address"] = address

    purpose = extract_by_patterns(
        material_text,
        [r"(?:мета сайту|ціль сайту|purpose)\s*[:\-]\s*(.+)"],
    )
    if purpose:
        result["site_purpose_description"] = purpose

    products = extract_by_patterns(
        material_text,
        [r"(?:товари|послуги|продукти|напрями діяльності)\s*[:\-]\s*(.+)"],
    )
    if products:
        result["main_products"] = products

    urls = re.findall(r"(?:https?://[^\s,]+)", material_text, re.I)
    if urls:
        result["liked_sites_list"] = sorted(set(urls))

    lower_text = material_text.lower()
    if any(token in lower_text for token in ["інтернет-магазин", "ecommerce", "кошик", "checkout"]):
        result["is_ecommerce"] = True
    if "не інтернет-магазин" in lower_text or "not ecommerce" in lower_text:
        result["is_ecommerce"] = False

    return normalize_by_schema(result, schema)


def try_parse_embedded_json(material_text: str, schema: dict[str, Any]) -> dict[str, Any]:
    default_payload = build_default_from_schema(schema)

    candidates = [material_text.strip()]
    brace_match = re.search(r"\{[\s\S]*\}", material_text)
    if brace_match:
        candidates.append(brace_match.group(0))

    for candidate in candidates:
        if not candidate:
            continue
        try:
            parsed = json.loads(candidate)
        except json.JSONDecodeError:
            continue

        if isinstance(parsed, dict):
            return normalize_by_schema(parsed, schema)

    enriched_payload = enrich_from_text(material_text, schema)
    if enriched_payload:
        return enriched_payload

    return default_payload


def generate_with_openai(material_text: str, placeholders: dict[str, Any], schema: dict[str, Any]) -> dict[str, Any]:
    api_key = os.getenv("OPENAI_API_KEY", "").strip()
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY is not set")

    if not OPENAI_AVAILABLE:
        raise RuntimeError("openai package is not installed")

    client = OpenAI(api_key=api_key, timeout=OPENAI_TIMEOUT_SECONDS, max_retries=OPENAI_MAX_RETRIES)
    user_prompt = get_generation_payload(material_text, placeholders, schema)

    response = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": user_prompt},
        ],
        response_format={"type": "json_object"},
        temperature=0,
    )

    content = response.choices[0].message.content or "{}"
    parsed = parse_json_from_response(content)
    if not isinstance(parsed, dict):
        raise RuntimeError("LLM повернула не JSON-об'єкт")

    return normalize_by_schema(parsed, schema)


def normalize_ecommerce_flag(value: Any) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, str):
        v = value.strip().lower()
        if v in {"true", "так", "yes", "1"}:
            return True
        if v in {"false", "ні", "no", "0", "уточнити"}:
            return False
    return False


def highlight_word_in_docx(file_path: Path, word: str = "уточнити") -> None:
    if not DOCX_AVAILABLE:
        return

    from docx.oxml import OxmlElement
    from docx.text.run import Run

    def copy_run_style(src_run: Any, dst_run: Any) -> None:
        dst_run.style = src_run.style
        src_font = src_run.font
        dst_font = dst_run.font
        dst_font.name = src_font.name
        dst_font.size = src_font.size
        dst_font.bold = src_font.bold
        dst_font.italic = src_font.italic
        dst_font.underline = src_font.underline
        if src_font.color is not None and src_font.color.rgb is not None:
            dst_font.color.rgb = src_font.color.rgb

    def process_paragraph(paragraph: Any) -> None:
        i = 0
        while i < len(paragraph.runs):
            run = paragraph.runs[i]
            pos = run.text.find(word)
            if pos != -1:
                text = run.text
                before = text[:pos]
                target = text[pos : pos + len(word)]
                after = text[pos + len(word) :]
                run.text = before
                r_el = OxmlElement("w:r")
                run._r.addnext(r_el)
                new_run_word = Run(r_el, paragraph)
                copy_run_style(run, new_run_word)
                new_run_word.text = target
                new_run_word.font.highlight_color = WD_COLOR_INDEX.YELLOW
                if after:
                    r_after = OxmlElement("w:r")
                    new_run_word._r.addnext(r_after)
                    new_run_after = Run(r_after, paragraph)
                    copy_run_style(run, new_run_after)
                    new_run_after.text = after
                i += 2
            else:
                i += 1

    docx = Document(file_path)
    for paragraph in docx.paragraphs:
        process_paragraph(paragraph)
    for table in docx.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)
    docx.save(file_path)


def render_docx(payload: dict[str, Any]) -> Path:
    if not DOCX_AVAILABLE:
        raise RuntimeError(
            "Пакети docxtpl/python-docx не встановлені. "
            f"Виконайте: {sys.executable} -m pip install docxtpl python-docx"
        )
    if not TEMPLATE_PATH.exists():
        raise RuntimeError("Не знайдено шаблон data/template.docx")

    doc = DocxTemplate(TEMPLATE_PATH)
    context = json.loads(json.dumps(payload, ensure_ascii=False))
    context["is_ecommerce"] = normalize_ecommerce_flag(context.get("is_ecommerce"))
    domain_value = context.get("domain")
    if RichText:
        rich_domain = RichText()
        if isinstance(domain_value, str) and domain_value.strip() and not is_placeholder_text(domain_value):
            domain_text = domain_value.strip()
            domain_url = domain_text if domain_text.startswith(("http://", "https://")) else f"https://{domain_text}"
            rich_domain.add(domain_text, url_id=doc.build_url_id(domain_url))
        else:
            rich_domain.add("уточнити")
        context["domain"] = rich_domain
    messengers = context.get("footer_messengers_list")
    if isinstance(messengers, list):
        formatted: list[str] = []
        for item in messengers:
            if isinstance(item, str) and item.strip():
                formatted.append(item.strip())
            elif isinstance(item, dict):
                parts = []
                for key in ("name", "value", "contact", "description"):
                    value = item.get(key)
                    if isinstance(value, str) and value.strip():
                        parts.append(value.strip())
                if parts:
                    formatted.append(" — ".join(parts[:2]))
        context["footer_messengers_list"] = ", ".join(formatted)

    output_name = f"tz_{uuid.uuid4().hex}.docx"
    output_path = Path(tempfile.gettempdir()) / output_name

    doc.render(context)
    doc.save(output_path)
    highlight_word_in_docx(output_path, word="уточнити")
    return output_path


def generate_payload(material_text: str, description_override: str = "") -> tuple[dict[str, Any], str, list[str]]:
    schema = load_json_file(SCHEMA_PATH)
    if not schema:
        raise RuntimeError("Не знайдено data/output_schema.json")

    placeholders = load_json_file(DESCRIPTION_PATH)
    if description_override:
        try:
            override_obj = json.loads(description_override)
            if isinstance(override_obj, dict):
                placeholders = override_obj
        except json.JSONDecodeError:
            pass

    warnings: list[str] = []
    try:
        payload = generate_with_openai(material_text, placeholders, schema)
        source = "openai"
    except Exception as exc:  # pragma: no cover - runtime path
        payload = try_parse_embedded_json(material_text, schema)
        source = "fallback"
        warnings.append(f"Fallback mode used: {safe_error_message(exc)}")

    payload = finalize_payload_for_template(payload, material_text, schema)
    return payload, source, warnings


def collect_input_material() -> str:
    text_input = (request.form.get("input_text") or "").strip()
    brief_text = read_uploaded_text(request.files.get("brief_file"))

    parts = [part.strip() for part in [text_input, brief_text] if part and part.strip()]
    return "\\n\\n".join(parts)


app = Flask(__name__, static_folder="web", static_url_path="")


@app.get("/")
def index() -> Any:
    return app.send_static_file("index.html")


@app.get("/api/health")
def health() -> Any:
    return jsonify({"status": "ok"})


@app.post("/api/generate")
def generate() -> Any:
    material_text = collect_input_material()
    if not material_text:
        return jsonify({"error": "Надайте бриф у тексті або завантажте файл."}), 400

    description_override = read_uploaded_text(request.files.get("description_file"))

    try:
        payload, source, warnings = generate_payload(material_text, description_override)
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500

    return jsonify(
        {
            "source": source,
            "warnings": warnings,
            "result_json": payload,
        }
    )


@app.post("/api/generate-docx")
def generate_docx() -> Any:
    started_at = time.perf_counter()
    material_text = collect_input_material()
    if not material_text:
        return jsonify({"error": "Надайте бриф у тексті або завантажте файл."}), 400

    description_override = read_uploaded_text(request.files.get("description_file"))
    try:
        payload, source, warnings = generate_payload(material_text, description_override)
        output_path = render_docx(payload)
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500

    @after_this_request
    def cleanup(response: Any) -> Any:
        try:
            output_path.unlink(missing_ok=True)
        except OSError:
            pass
        return response

    filename = f"TZ_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    response = send_file(
        output_path,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response.headers["X-Generator-Source"] = source
    response.headers["X-Generator-Warning-Count"] = str(len(warnings))
    if warnings:
        warning_header = warnings[0].replace("\n", " ").strip()
        warning_header = warning_header.encode("latin-1", errors="replace").decode("latin-1")
        response.headers["X-Generator-Warning"] = warning_header[:500]
    response.headers["X-Generator-Duration-Seconds"] = f"{time.perf_counter() - started_at:.2f}"
    return response


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.getenv("PORT", "8000")), debug=True)
